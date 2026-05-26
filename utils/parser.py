"""
utils/parser.py — Resume parsing engine
=========================================
Extracts raw text from PDF and DOCX resume files.
Supports:
  - pdfplumber for PDF (handles multi-column, tables, headers)
  - python-docx for DOCX
  - basic heuristics to detect education and certification sections

Public API
----------
    parse_resume(filepath: str) -> dict
        Returns {"text": str, "education": str, "certifications": str, "error": str|None}
"""

import re
import os
import logging
from pathlib import Path

logger = logging.getLogger(__name__)


# ── Section header patterns ───────────────────────────────────────────
_EDUCATION_PATTERNS = re.compile(
    r"(education|academic|qualifications?|degree|university|college|school)",
    re.IGNORECASE,
)
_CERT_PATTERNS = re.compile(
    r"(certif|certificat|aws certified|google cloud|microsoft certified|"
    r"pmp|cissp|ceh|ccna|comptia|oracle|ibm|coursera|udemy|certificate)",
    re.IGNORECASE,
)


def parse_resume(filepath: str) -> dict:
    """
    Parse a resume file and return structured text.

    Parameters
    ----------
    filepath : str
        Absolute path to the uploaded resume (PDF or DOCX).

    Returns
    -------
    dict with keys:
        text          — full extracted text (str)
        education     — education-related lines (str)
        certifications — certification lines (str)
        error         — error message if parsing failed (str | None)
    """
    ext = Path(filepath).suffix.lower()

    try:
        if ext == ".pdf":
            raw_text = _parse_pdf(filepath)
        elif ext in (".docx", ".doc"):
            raw_text = _parse_docx(filepath)
        else:
            return {
                "text": "",
                "education": "",
                "certifications": "",
                "error": f"Unsupported file type: {ext}",
            }
    except Exception as exc:
        logger.error("Resume parsing failed for %s: %s", filepath, exc, exc_info=True)
        return {
            "text": "",
            "education": "",
            "certifications": "",
            "error": f"Failed to parse resume: {exc}",
        }

    if not raw_text.strip():
        return {
            "text": "",
            "education": "",
            "certifications": "",
            "error": "Resume appears to be empty or contains only images.",
        }

    education, certifications = _extract_sections(raw_text)

    return {
        "text": raw_text,
        "education": education,
        "certifications": certifications,
        "error": None,
    }


# ── PDF parser ────────────────────────────────────────────────────────
def _parse_pdf(filepath: str) -> str:
    """Extract text from a PDF using pdfplumber."""
    import pdfplumber

    pages_text = []
    with pdfplumber.open(filepath) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                pages_text.append(text)

    return "\n".join(pages_text)


# ── DOCX parser ───────────────────────────────────────────────────────
def _parse_docx(filepath: str) -> str:
    """Extract text from a DOCX using python-docx (paragraphs + tables)."""
    from docx import Document

    doc = Document(filepath)
    chunks: list[str] = []

    # ── Paragraphs ────────────────────────────────────────────────────
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            chunks.append(text)

    # ── Tables ────────────────────────────────────────────────────────
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(
                cell.text.strip() for cell in row.cells if cell.text.strip()
            )
            if row_text:
                chunks.append(row_text)

    return "\n".join(chunks)


# ── Section heuristics ────────────────────────────────────────────────
def _extract_sections(text: str) -> tuple[str, str]:
    """
    Walk the lines of *text* and collect education/certification snippets.
    Returns (education_str, certifications_str).
    """
    lines = text.split("\n")
    education_lines: list[str] = []
    cert_lines: list[str] = []

    _in_edu_section  = False
    _in_cert_section = False

    for line in lines:
        stripped = line.strip()
        if not stripped:
            _in_edu_section  = False
            _in_cert_section = False
            continue

        # Detect section headers
        if _EDUCATION_PATTERNS.search(stripped) and len(stripped) < 60:
            _in_edu_section  = True
            _in_cert_section = False
            education_lines.append(stripped)
            continue

        if _CERT_PATTERNS.search(stripped) and len(stripped) < 80:
            _in_cert_section = True
            _in_edu_section  = False
            cert_lines.append(stripped)
            continue

        if _in_edu_section:
            education_lines.append(stripped)
        elif _in_cert_section:
            cert_lines.append(stripped)

    return "\n".join(education_lines[:20]), "\n".join(cert_lines[:20])
